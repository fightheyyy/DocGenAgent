"""
ChromaDB RAG工具 - 重新设计版
核心功能：文档embedding处理、智能搜索、基于模板字段的内容填充
"""
import os
import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import hashlib

try:
    import chromadb
    from chromadb.config import Settings
    import fitz  # PyMuPDF for PDF
    from docx import Document as DocxDocument
except ImportError as e:
    print(f"警告: RAG工具依赖未安装: {e}")
    print("请安装: pip install chromadb PyMuPDF python-docx")

try:
    from src.base_tool import Tool
    from src.pdf_embedding_service import PDFEmbeddingService
    PDF_EMBEDDING_SERVICE_AVAILABLE = True
except ImportError:
    # 如果无法导入Tool，创建一个基础类
    class Tool:
        def __init__(self):
            self.name = "base_tool"
            self.description = "基础工具类"
        
        def execute(self, action: str, **kwargs) -> str:
            return "基础工具执行"
    
    PDF_EMBEDDING_SERVICE_AVAILABLE = False

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentExtractor:
    """文档内容提取器"""
    
    def extract_content(self, file_path: str) -> str:
        """提取文档内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._extract_from_text(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            raise RuntimeError(f"文档内容提取失败: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """从PDF提取内容"""
        content = ""
        doc = fitz.open(file_path)
        for page in doc:
            content += page.get_text()
        doc.close()
        return content.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """从DOCX提取内容"""
        doc = DocxDocument(file_path)
        content = []
        
        # 提取段落内容
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content.append(f"表格行: {row_text}")
        
        return "\n".join(content)
    
    def _extract_from_text(self, file_path: str) -> str:
        """从文本文件提取内容"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()

class ChromaVectorStore:
    """ChromaDB向量存储"""
    
    def __init__(self, storage_dir: str):
        self.storage_dir = storage_dir
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                self.client = chromadb.PersistentClient(
                    path=storage_dir,
                    settings=Settings(
                        allow_reset=True, 
                        anonymized_telemetry=False,
                        is_persistent=True
                    )
                )
                self.collection = self.client.get_or_create_collection(
                    name="documents",
                    metadata={"hnsw:space": "cosine"}
                )
                logger.info(f"✅ ChromaVectorStore初始化成功: {storage_dir}")
                return
                
            except Exception as e:
                error_msg = str(e)
                if "already exists" in error_msg and "different settings" in error_msg:
                    logger.warning(f"⚠️ ChromaDB实例冲突，尝试重置... (尝试 {attempt + 1}/{max_retries})")
                    
                    # 尝试重置ChromaDB连接
                    try:
                        if hasattr(self, 'client') and self.client:
                            self.client.reset()
                    except:
                        pass
                    
                    # 等待一下再重试
                    import time
                    time.sleep(1)
                    
                    if attempt == max_retries - 1:
                        # 最后一次尝试：删除并重新创建
                        try:
                            import shutil
                            if os.path.exists(storage_dir):
                                logger.info(f"🔄 清理ChromaDB目录: {storage_dir}")
                                shutil.rmtree(storage_dir)
                                os.makedirs(storage_dir, exist_ok=True)
                        except Exception as cleanup_error:
                            logger.warning(f"⚠️ 清理失败: {cleanup_error}")
                else:
                    logger.error(f"❌ ChromaVectorStore初始化失败: {e}")
                    if attempt == max_retries - 1:
                        raise
                    else:
                        import time
                        time.sleep(1)
        
        # 如果所有重试都失败了
        raise RuntimeError("ChromaVectorStore初始化失败，已达到最大重试次数")
    
    def add_document(self, doc_id: str, content: str, metadata: Dict[str, Any]):
        """添加文档到向量库"""
        # 将长文档分块
        chunks = self._split_content(content)
        
        ids = []
        documents = []
        metadatas = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            ids.append(chunk_id)
            documents.append(chunk)
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
            metadatas.append(chunk_metadata)
        
        self.collection.add(
            ids=ids,
            documents=documents,
            metadatas=metadatas
        )
        
        return len(chunks)
    
    def search_documents(self, query: str, n_results: int = 5, where_filter: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """搜索相关文档，支持元数据过滤"""
        query_params = {
            "query_texts": [query],
            "n_results": n_results
        }
        if where_filter:
            query_params["where"] = where_filter
            logger.info(f"🔍 使用元数据过滤器进行搜索: {where_filter}")
            
        results = self.collection.query(**query_params)
        
        formatted_results = []
        if results['documents'] and len(results['documents']) > 0:
            for i in range(len(results['documents'][0])):
                formatted_results.append({
                    'content': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i]
                })
        
        return formatted_results
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """获取所有文档信息"""
        try:
            results = self.collection.get()
            documents = []
            
            # 按文档ID分组
            doc_groups = {}
            for i, doc_id in enumerate(results['ids']):
                base_id = doc_id.split('_chunk_')[0]
                if base_id not in doc_groups:
                    doc_groups[base_id] = {
                        'id': base_id,
                        'metadata': results['metadatas'][i],
                        'chunks': 0
                    }
                doc_groups[base_id]['chunks'] += 1
            
            return list(doc_groups.values())
        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return []
    
    def clear_all(self):
        """清空所有文档"""
        self.client.delete_collection("documents")
        self.collection = self.client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def _split_content(self, content: str, chunk_size: int = 1000) -> List[str]:
        """将内容分块"""
        if len(content) <= chunk_size:
            return [content]
        
        chunks = []
        sentences = content.split('。')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk + sentence + '。') <= chunk_size:
                current_chunk += sentence + '。'
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + '。'
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [content]

class TemplateFieldProcessor:
    """模板字段处理器 - 核心智能填充功能"""
    
    def __init__(self, deepseek_client=None):
        self.deepseek_client = deepseek_client
    
    def fill_template_fields(self, template_fields_json: Dict[str, str], 
                           vector_store: ChromaVectorStore) -> Dict[str, str]:
        """
        基于模板字段JSON进行智能搜索和内容填充
        
        Args:
            template_fields_json: 模板字段JSON，格式为 {"字段名": "字段描述或要求"}
            vector_store: 向量存储实例
            
        Returns:
            填充好的字段JSON，格式为 {"字段名": "填充的具体内容"}
        """
        logger.info(f"🔍 开始基于模板字段进行智能填充，共 {len(template_fields_json)} 个字段")
        
        filled_fields = {}
        
        for field_name, field_requirement in template_fields_json.items():
            logger.info(f"📝 处理字段: {field_name}")
            
            # 1. 基于字段要求搜索相关内容
            search_results = vector_store.search_documents(
                query=f"{field_name} {field_requirement}",
                n_results=3
            )
            
            # 2. 提取搜索到的内容
            relevant_content = []
            for result in search_results:
                relevant_content.append(result['content'])
            
            # 3. 使用AI生成字段内容（如果有AI客户端）
            if self.deepseek_client and relevant_content:
                filled_content = self._generate_field_content_with_ai(
                    field_name, field_requirement, relevant_content
                )
            else:
                # 基础模式：直接使用搜索到的最相关内容
                filled_content = self._generate_field_content_basic(
                    field_name, field_requirement, relevant_content
                )
            
            filled_fields[field_name] = filled_content
            logger.info(f"✅ 字段 {field_name} 填充完成，内容长度: {len(filled_content)} 字符")
        
        logger.info(f"🎉 所有字段填充完成！")
        return filled_fields
    
    def _generate_field_content_with_ai(self, field_name: str, field_requirement: str, 
                                       relevant_content: List[str]) -> str:
        """使用AI生成字段内容"""
        content_text = "\n".join(relevant_content)
        
        prompt = f"""
你是一个专业的文档处理助手。请根据以下信息为字段生成合适的内容：

字段名称：{field_name}
字段要求：{field_requirement}

相关资料内容：
{content_text}

任务要求：
1. 基于相关资料内容，为该字段生成专业、准确的内容
2. 内容应该符合字段要求和描述
3. 保持内容的专业性和完整性
4. 如果资料内容不足，请基于字段要求进行合理补充
5. 内容长度适中，重点突出

请直接返回该字段的具体内容，不要包含解释文字。
"""
        
        try:
            response = self.deepseek_client.chat([{"role": "user", "content": prompt}])
            return response.strip() if response else self._generate_field_content_basic(
                field_name, field_requirement, relevant_content
            )
        except Exception as e:
            logger.warning(f"AI生成失败，使用基础模式: {e}")
            return self._generate_field_content_basic(
                field_name, field_requirement, relevant_content
            )
    
    def _generate_field_content_basic(self, field_name: str, field_requirement: str, 
                                     relevant_content: List[str]) -> str:
        """基础模式生成字段内容"""
        if not relevant_content:
            return f"[{field_name}]：{field_requirement}（待补充具体内容）"
        
        # 选择最相关的内容作为基础
        base_content = relevant_content[0]
        
        # 简单的内容处理
        if len(base_content) > 200:
            # 截取前200字符作为摘要
            summary = base_content[:200] + "..."
            return f"{summary}\n\n（基于相关资料整理，如需详细信息请参考原始文档）"
        else:
            return base_content

class RAGTool(Tool):
    """重新设计的RAG工具 - 核心三功能"""
    
    def __init__(self, storage_dir: str = "rag_storage", deepseek_client=None):
        super().__init__()
        self.name = "rag_tool"
        self.description = """智能文档检索工具 - 支持文本、图片、表格的精确搜索

🔍 **支持的操作 (action):**

1. **search** - 通用搜索（返回混合内容）
   参数: {"action": "search", "query": "关键词", "top_k": 5}

2. **search_images** - 专门搜索图片 🖼️
   参数: {"action": "search_images", "query": "关键词", "top_k": 5}
   示例: {"action": "search_images", "query": "医灵古庙", "top_k": 8}

3. **search_tables** - 专门搜索表格 📊
   参数: {"action": "search_tables", "query": "关键词", "top_k": 5}
   示例: {"action": "search_tables", "query": "影响评估", "top_k": 5}

4. **count_images** - 统计图片数量 📈
   参数: {"action": "count_images", "query": "关键词"}
   示例: {"action": "count_images", "query": "医灵古庙"}

5. **count_tables** - 统计表格数量 📈
   参数: {"action": "count_tables", "query": "关键词"}
   示例: {"action": "count_tables", "query": "评估标准"}

6. **list** - 列出所有文档
   参数: {"action": "list"}

7. **process_parsed_folder** - 处理解析文件夹
   参数: {"action": "process_parsed_folder", "folder_path": "路径"}

⚠️ **重要提示:**
- 搜索图片请使用 search_images，不要使用 search + search_type
- 统计数量请使用 count_images/count_tables，不要使用 limit 参数
- top_k 参数控制返回结果数量（默认5）
- 所有操作都需要明确指定 action 参数

💡 **使用场景:**
- 问"有多少张图片？" → 使用 count_images
- 问"检索N张图片" → 使用 search_images + top_k
- 问"搜索表格数据" → 使用 search_tables
"""
        
        self.storage_dir = storage_dir
        os.makedirs(storage_dir, exist_ok=True)
        
        # 初始化组件
        self.extractor = DocumentExtractor()
        self.vector_store = ChromaVectorStore(storage_dir)
        self.field_processor = TemplateFieldProcessor(deepseek_client)
        
        # 初始化改进的PDF embedding服务
        self.pdf_embedding_service = None
        if PDF_EMBEDDING_SERVICE_AVAILABLE:
            try:
                self.pdf_embedding_service = PDFEmbeddingService(
                    chroma_db_path=storage_dir,
                    collection_name="documents"
                )
                logger.info("✅ PDF Embedding Service integrated successfully")
            except Exception as e:
                logger.warning(f"⚠️ PDF Embedding Service initialization failed: {e}")
        else:
            logger.warning("⚠️ PDF Embedding Service not available")
    
    def execute(self, action: str, **kwargs) -> str:
        """执行RAG操作"""
        try:
            if action == "upload":
                return self._upload_document(**kwargs)
            elif action == "upload_image":
                return self._upload_image(**kwargs)
            elif action == "search":
                return self._search_documents(**kwargs)
            elif action == "search_images":
                return self._search_images(**kwargs)
            elif action == "search_tables":
                return self._search_tables(**kwargs)
            elif action == "count_images":
                return self._count_images(**kwargs)
            elif action == "count_tables":
                return self._count_tables(**kwargs)
            elif action == "fill_fields":
                return self._fill_template_fields(**kwargs)
            elif action == "list":
                return self._list_documents()
            elif action == "clear":
                return self._clear_documents()
            elif action == "process_parsed_folder":
                folder_path = kwargs.get("folder_path")
                project_name = kwargs.get("project_name", "")
                if not folder_path:
                    return "❌ 请提供解析文件夹路径 (folder_path参数)"
                return self._process_parsed_folder(folder_path, project_name)
            else:
                return f"""❌ 不支持的操作: {action}

📋 **支持的操作列表:**
• upload - 上传文档
• upload_image - 上传图片并生成AI描述
• search - 通用搜索
• search_images - 搜索图片
• search_tables - 搜索表格  
• count_images - 统计图片数量
• count_tables - 统计表格数量
• list - 列出文档
• process_parsed_folder - 处理解析文件夹

💡 **使用示例:**
• 上传文档: {{"action": "upload", "file_path": "document.pdf"}}
• 上传图片: {{"action": "upload_image", "image_path": "image.jpg", "description": "可选描述"}}
• 搜索图片: {{"action": "search_images", "query": "医灵古庙", "top_k": 8}}
• 统计图片: {{"action": "count_images", "query": "医灵古庙"}}
"""
        
        except Exception as e:
            error_msg = str(e)
            
            # 检查是否是参数错误
            if "unexpected keyword argument" in error_msg:
                if "search_type" in error_msg:
                    return """❌ 参数错误: search操作不支持search_type参数

✅ **正确做法:**
• 搜索图片: {"action": "search_images", "query": "关键词"}
• 搜索表格: {"action": "search_tables", "query": "关键词"}
• 统计图片: {"action": "count_images", "query": "关键词"}
"""
                elif "limit" in error_msg:
                    return """❌ 参数错误: 不支持limit参数

✅ **正确做法:**
• 使用top_k参数: {"action": "search_images", "query": "关键词", "top_k": 8}
• 或使用默认值: {"action": "search_images", "query": "关键词"}
"""
            
            logger.error(f"RAG操作失败: {e}")
            return f"❌ 操作失败: {str(e)}"
    
    def _upload_document(self, file_path: str, filename: str = "", metadata: Optional[Dict[str, Any]] = None) -> str:
        """
        上传并处理单个文档
        
        Args:
            file_path: 文档文件路径
            filename: 文档在系统中的名称（可选）
            metadata: 附加的元数据（可选）
            
        Returns:
            处理结果
        """
        try:
            doc_name = filename if filename else os.path.basename(file_path)
            doc_id = hashlib.md5(doc_name.encode()).hexdigest()
            
            logger.info(f"📤 开始处理文档: {doc_name}")
            
            # 1. 提取内容
            content = self.extractor.extract_content(file_path)
            
            # 2. 准备元数据
            if metadata is None:
                metadata = {}
            
            # 确保基本元数据存在
            metadata.setdefault("source", doc_name)
            metadata.setdefault("upload_time", datetime.now().isoformat())
            metadata.setdefault("file_size", os.path.getsize(file_path))

            # 3. 添加到向量库
            chunks_count = self.vector_store.add_document(doc_id, content, metadata)
            
            logger.info(f"✅ 文档处理完成: {doc_name} (ID: {doc_id}), 共 {chunks_count} 个块")
            
            return f"✅ 文档 '{doc_name}' 上传并处理成功，共分为 {chunks_count} 个内容块。"
        except FileNotFoundError:
            return f"❌ 文件未找到: {file_path}"
        except Exception as e:
            return f"❌ 文档处理失败: {str(e)}"
    
    def _upload_image(self, image_path: str, description: str = "") -> str:
        """
        上传图片，生成AI描述，并将其嵌入到统一知识库
        
        Args:
            image_path: 图片文件路径
            description: 用户提供的可选描述
            
        Returns:
            处理结果的JSON字符串
        """
        try:
            from pathlib import Path
            
            if not image_path or not os.path.exists(image_path):
                return json.dumps({"status": "error", "message": "图片路径不存在或未提供"}, ensure_ascii=False)
            
            # 导入OpenRouter客户端用于AI描述生成
            try:
                from src.openrouter_client import OpenRouterClient
                openrouter_client = OpenRouterClient()
            except Exception as e:
                return json.dumps({"status": "error", "message": f"AI服务初始化失败: {e}"}, ensure_ascii=False)
            
            logger.info(f"🚀 开始处理上传的图片: {image_path}")
            
            # 1. 使用Gemini生成图片描述
            prompt = "请详细描述这张图片的内容，包括场景、物体、人物、风格和任何可见的文本。"
            ai_description = openrouter_client.get_image_description_gemini(image_path, prompt=prompt)
            
            if "Error:" in ai_description:
                raise Exception(f"AI描述生成失败: {ai_description}")

            final_description = f"用户描述: {description}\n\nAI描述: {ai_description}" if description else ai_description
            logger.info(f"📝 生成的描述: {final_description[:150]}...")

            # 2. 准备元数据并嵌入到ChromaDB
            image_name = Path(image_path).name
            doc_id = hashlib.md5(image_name.encode()).hexdigest()
            
            metadata = {
                "source": image_name,
                "document_type": "Image",
                "upload_time": datetime.now().isoformat(),
                "user_provided_description": bool(description),
                "file_size": os.path.getsize(image_path)
            }
            
            # 3. 添加到向量库
            chunks_count = self.vector_store.add_document(doc_id, final_description, metadata)
            
            logger.info(f"✅ 图片处理完成: {image_name} (ID: {doc_id}), 共 {chunks_count} 个块")
            
            result = {
                "status": "success",
                "message": "图片处理和嵌入成功",
                "image_source": image_path,
                "chunks_count": chunks_count,
                "generated_description": final_description
            }
            
            return json.dumps(result, ensure_ascii=False, indent=2)

        except Exception as e:
            logger.error(f"❌ 图片上传和嵌入流程失败: {e}")
            return json.dumps({"status": "error", "message": str(e)}, ensure_ascii=False)
    
    def _search_documents(self, query: str, top_k: int = 5, metadata_filter: Optional[Dict[str, Any]] = None, project_name: Optional[str] = None) -> str:
        """
        根据查询文本搜索文档 - 支持智能项目隔离
        
        Args:
            query (str): 搜索的关键词或问题。
            top_k (int): 返回最相关的结果数量，默认为5。
            metadata_filter (dict, optional): 用于精确过滤的元数据。
            project_name (str, optional): 项目名称过滤器（可选，如果不提供会尝试自动提取）
                *   示例1 (只搜图片): `{"document_type": "Image"}`
                *   示例2 (只搜文本): `{"document_type": "Text"}`
                *   示例3 (只搜特定文件): `{"source_document": "your_file.pdf"}`
                *   示例4 (复合查询：只搜特定文件中的图片): `{"$and": [{"document_type": "Image"}, {"source_document": "your_file.pdf"}]}`
            
        Returns:
            str: 包含搜索结果的JSON字符串
        """
        logger.info(f"执行文档搜索: query='{query}', top_k={top_k}, filter={metadata_filter}")
        
        try:
            # 🆕 智能项目名称提取
            if not project_name:
                project_name = self._extract_project_name_from_query(query)
                if project_name:
                    logger.info(f"🎯 自动提取项目名称: {project_name}")
            
            # 🆕 如果有项目名称且PDF Embedding Service可用，优先使用项目隔离搜索
            if project_name and self.pdf_embedding_service:
                logger.info(f"🔍 使用项目隔离搜索: {project_name}")
                results = self.pdf_embedding_service.search(
                    query=query,
                    top_k=top_k,
                    content_type=None,  # 搜索所有类型
                    source_file_filter=None,
                    project_name=project_name
                )
                
                simplified_results = []
                for res in results:
                    simplified_results.append({
                        "content": res.get("content", ""),
                        "metadata": res.get("metadata", {}),
                        "distance": res.get("distance", 0.0)
                    })
                
                result_data = {
                    "status": "success", 
                    "results": simplified_results,
                    "total_count": len(simplified_results),
                    "search_method": "pdf_embedding_service_with_project_isolation"
                }
                
                result_data["project_isolation"] = {
                    "enabled": True,
                    "project_name": project_name,
                    "message": f"🔒 已限制搜索范围至项目: {project_name}"
                }
                
                return json.dumps(result_data, ensure_ascii=False)
            
            # 🔄 回退到传统向量搜索
            else:
                logger.info("🔍 使用传统向量搜索")
                results = self.vector_store.search_documents(
                    query=query, 
                    n_results=top_k,
                    where_filter=metadata_filter
                )
                
                # 简化输出
                simplified_results = []
                for res in results:
                    simplified_results.append({
                        "content": res.get("content"),
                        "metadata": res.get("metadata"),
                        "distance": res.get("distance")
                    })

                result_data = {
                    "status": "success", 
                    "results": simplified_results,
                    "total_count": len(simplified_results),
                    "search_method": "vector_store"
                }
                
                if project_name:
                    result_data["project_isolation"] = {
                        "enabled": False,
                        "attempted_project": project_name,
                        "message": f"⚠️ PDF Embedding Service不可用，使用传统搜索（无项目隔离）"
                    }
                else:
                    result_data["project_isolation"] = {
                        "enabled": False,
                        "message": "⚠️ 未启用项目隔离，搜索了所有项目"
                    }

                return json.dumps(result_data, ensure_ascii=False)
                
        except Exception as e:
            logger.error(f"文档搜索失败: {e}")
            return json.dumps({"status": "error", "message": str(e)})
    
    def _search_images(self, query: str, top_k: int = 5, source_file_filter: Optional[str] = None, project_name: Optional[str] = None) -> str:
        """
        搜索图片内容 - 支持智能项目隔离
        
        Args:
            query: 搜索关键词
            top_k: 返回结果数量
            source_file_filter: 源文件过滤器
            project_name: 项目名称过滤器（可选，如果不提供会尝试自动提取）
            
        Returns:
            包含搜索结果的JSON字符串
        """
        if not self.pdf_embedding_service:
            return json.dumps({
                "status": "error", 
                "message": "PDF Embedding Service not available"
            })
        
        try:
            # 🆕 智能项目名称提取
            if not project_name:
                # 尝试从查询中提取项目名称
                project_name = self._extract_project_name_from_query(query)
                if project_name:
                    logger.info(f"🎯 自动提取项目名称: {project_name}")
            
            results = self.pdf_embedding_service.search_images_only(
                query=query, 
                top_k=top_k, 
                source_file_filter=source_file_filter,
                project_name=project_name  # 🆕 项目隔离参数
            )
            
            simplified_results = []
            for res in results:
                simplified_results.append({
                    "content": res.get("content", ""),
                    "metadata": res.get("metadata", {}),
                    "content_type": res.get("content_type", "image"),
                    "distance": res.get("distance", 0.0)
                })
            
            # 🆕 添加项目隔离信息到返回结果
            result_data = {
                "status": "success", 
                "results": simplified_results,
                "total_count": len(simplified_results)
            }
            
            if project_name:
                result_data["project_isolation"] = {
                    "enabled": True,
                    "project_name": project_name,
                    "message": f"🔒 已限制搜索范围至项目: {project_name}"
                }
            else:
                result_data["project_isolation"] = {
                    "enabled": False,
                    "message": "⚠️ 未启用项目隔离，搜索了所有项目"
                }
            
            return json.dumps(result_data, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"图片搜索失败: {e}")
            return json.dumps({"status": "error", "message": str(e)})
    
    def _extract_project_name_from_query(self, query: str) -> Optional[str]:
        """
        从查询中智能提取项目名称
        
        Args:
            query: 搜索查询
            
        Returns:
            提取到的项目名称，如果未找到则返回None
        """
        try:
            # 获取所有可用项目
            if hasattr(self.pdf_embedding_service, 'get_available_projects'):
                available_projects = self.pdf_embedding_service.get_available_projects()
                
                # 按长度排序，优先匹配更长的项目名称
                available_projects = sorted(available_projects, key=len, reverse=True)
                
                # 🔍 正确的匹配逻辑：检查查询是否包含在项目名称中
                for project in available_projects:
                    if query in project:
                        return project
                
                # 🔍 智能关键词匹配：提取项目的核心关键词
                for project in available_projects:
                    # 移除常见后缀，提取核心关键词
                    project_core = project.replace("设计方案", "").replace("修缮设计方案", "").replace("项目", "").replace("文物", "").strip()
                    
                    # 检查查询是否包含核心关键词
                    if project_core and query in project_core:
                        return project
                    
                    # 检查核心关键词是否包含在查询中
                    if project_core and project_core in query:
                        return project
                
                # 🔍 分词匹配：处理复合词情况
                for project in available_projects:
                    # 分解项目名称为关键词列表
                    project_keywords = []
                    
                    # 提取主要关键词
                    if "宗祠" in project:
                        project_keywords.extend(["宗祠"])
                        # 提取宗祠前的姓氏
                        if "氏宗祠" in project:
                            idx = project.find("氏宗祠")
                            if idx > 0:
                                surname = project[idx-1:idx+1]  # 如"刘氏"
                                project_keywords.append(surname)
                    
                    if "古庙" in project:
                        project_keywords.extend(["古庙"])
                        # 提取古庙前的名称
                        idx = project.find("古庙")
                        if idx > 0:
                            # 尝试提取2-3个字符的名称
                            for length in [3, 2]:
                                if idx >= length:
                                    name = project[idx-length:idx]
                                    if name not in ["设计", "修缮", "方案"]:
                                        project_keywords.append(name)
                                        break
                    
                    # 检查是否有关键词匹配
                    for keyword in project_keywords:
                        if keyword in query:
                            return project
            
            return None
            
        except Exception as e:
            logger.warning(f"项目名称提取失败: {e}")
            return None
    
    def _search_tables(self, query: str, top_k: int = 5, source_file_filter: Optional[str] = None, project_name: Optional[str] = None) -> str:
        """
        搜索表格内容 - 支持智能项目隔离
        
        Args:
            query: 搜索关键词
            top_k: 返回结果数量
            source_file_filter: 源文件过滤器
            project_name: 项目名称过滤器（可选，如果不提供会尝试自动提取）
            
        Returns:
            包含搜索结果的JSON字符串
        """
        if not self.pdf_embedding_service:
            return json.dumps({
                "status": "error", 
                "message": "PDF Embedding Service not available"
            })
        
        try:
            # 🆕 智能项目名称提取
            if not project_name:
                project_name = self._extract_project_name_from_query(query)
                if project_name:
                    logger.info(f"🎯 自动提取项目名称: {project_name}")
            
            results = self.pdf_embedding_service.search_tables_only(
                query=query, 
                top_k=top_k, 
                source_file_filter=source_file_filter,
                project_name=project_name  # 🆕 项目隔离参数
            )
            
            simplified_results = []
            for res in results:
                simplified_results.append({
                    "content": res.get("content", ""),
                    "metadata": res.get("metadata", {}),
                    "content_type": res.get("content_type", "table"),
                    "distance": res.get("distance", 0.0)
                })
            
            # 🆕 添加项目隔离信息到返回结果
            result_data = {
                "status": "success", 
                "results": simplified_results,
                "total_count": len(simplified_results)
            }
            
            if project_name:
                result_data["project_isolation"] = {
                    "enabled": True,
                    "project_name": project_name,
                    "message": f"🔒 已限制搜索范围至项目: {project_name}"
                }
            else:
                result_data["project_isolation"] = {
                    "enabled": False,
                    "message": "⚠️ 未启用项目隔离，搜索了所有项目"
                }
            
            return json.dumps(result_data, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"表格搜索失败: {e}")
            return json.dumps({"status": "error", "message": str(e)})
    
    def _count_images(self, query: str = "", source_file_filter: Optional[str] = None) -> str:
        """
        统计图片数量
        
        Args:
            query: 搜索关键词（可选）
            source_file_filter: 源文件过滤器
            
        Returns:
            包含统计结果的JSON字符串
        """
        if not self.pdf_embedding_service:
            return json.dumps({
                "status": "error", 
                "message": "PDF Embedding Service not available"
            })
        
        try:
            # 如果有查询词，进行搜索
            if query:
                results = self.pdf_embedding_service.search_images_only(
                    query=query, 
                    top_k=100,  # 获取更多结果来统计
                    source_file_filter=source_file_filter
                )
                count = len(results)
                message = f"找到 {count} 张包含'{query}'的图片"
            else:
                # 获取所有图片统计
                stats = self.pdf_embedding_service.get_collection_stats()
                count = stats.get("image_embeddings", 0)
                message = f"系统中共有 {count} 张图片"
            
            return json.dumps({
                "status": "success", 
                "count": count,
                "message": message,
                "content_type": "image"
            }, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"图片统计失败: {e}")
            return json.dumps({"status": "error", "message": str(e)})
    
    def _count_tables(self, query: str = "", source_file_filter: Optional[str] = None) -> str:
        """
        统计表格数量
        
        Args:
            query: 搜索关键词（可选）
            source_file_filter: 源文件过滤器
            
        Returns:
            包含统计结果的JSON字符串
        """
        if not self.pdf_embedding_service:
            return json.dumps({
                "status": "error", 
                "message": "PDF Embedding Service not available"
            })
        
        try:
            # 如果有查询词，进行搜索
            if query:
                results = self.pdf_embedding_service.search_tables_only(
                    query=query, 
                    top_k=100,  # 获取更多结果来统计
                    source_file_filter=source_file_filter
                )
                count = len(results)
                message = f"找到 {count} 个包含'{query}'的表格"
            else:
                # 获取所有表格统计
                stats = self.pdf_embedding_service.get_collection_stats()
                count = stats.get("table_embeddings", 0)
                message = f"系统中共有 {count} 个表格"
            
            return json.dumps({
                "status": "success", 
                "count": count,
                "message": message,
                "content_type": "table"
            }, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"表格统计失败: {e}")
            return json.dumps({"status": "error", "message": str(e)})
    
    def _fill_template_fields(self, template_fields_json: Dict[str, str]) -> str:
        """基于模板字段JSON进行智能填充 - 核心功能"""
        try:
            logger.info("🎯 开始模板字段智能填充")
            
            # 验证输入
            if not isinstance(template_fields_json, dict):
                return "❌ template_fields_json 必须是字典格式"
            
            if not template_fields_json:
                return "❌ template_fields_json 不能为空"
            
            # 执行智能填充
            filled_fields = self.field_processor.fill_template_fields(
                template_fields_json, self.vector_store
            )
            
            # 格式化返回结果
            result = f"✅ 模板字段智能填充完成！\n\n"
            result += f"📋 输入字段: {len(template_fields_json)} 个\n"
            result += f"📝 填充字段: {len(filled_fields)} 个\n\n"
            result += "📄 填充结果:\n"
            result += "=" * 50 + "\n"
            
            for field_name, filled_content in filled_fields.items():
                result += f"🔸 {field_name}:\n"
                result += f"   {filled_content[:100]}{'...' if len(filled_content) > 100 else ''}\n\n"
            
            result += "=" * 50 + "\n"
            result += f"💾 完整填充结果JSON:\n{json.dumps(filled_fields, ensure_ascii=False, indent=2)}"
            
            return result
            
        except Exception as e:
            logger.error(f"模板字段填充失败: {e}")
            return f"❌ 模板字段填充失败: {str(e)}"
    
    def _list_documents(self) -> str:
        """列出所有已上传的文档"""
        try:
            documents = self.vector_store.get_all_documents()
            
            if not documents:
                return "📚 当前没有已上传的文档"
            
            result = f"📚 已上传文档列表 (共 {len(documents)} 个):\n\n"
            
            for i, doc in enumerate(documents, 1):
                result += f"📄 文档 {i}:\n"
                result += f"   📁 文件名: {doc['metadata'].get('filename', '未知')}\n"
                result += f"   🆔 文档ID: {doc['id']}\n"
                result += f"   📊 分块数: {doc['chunks']} 个\n"
                result += f"   ⏰ 上传时间: {doc['metadata'].get('upload_time', '未知')}\n\n"
            
            return result
            
        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return f"❌ 获取文档列表失败: {str(e)}"
    
    def _clear_documents(self) -> str:
        """清空所有文档"""
        try:
            self.vector_store.clear_all()
            return "✅ 所有文档已成功清空。"
        except Exception as e:
            return f"❌ 清空文档失败: {str(e)}"
    
    def _process_parsed_folder(self, folder_path: str, project_name: str = "") -> str:
        """
        处理PDF解析工具生成的文件夹，将其中的文本内容添加到RAG知识库
        """
        try:
            content_file = os.path.join(folder_path, "parsed_content.json")
            if not os.path.exists(content_file):
                return f"❌ 'parsed_content.json' not found in {folder_path}"
                
            with open(content_file, 'r', encoding='utf-8') as f:
                parsed_data = json.load(f)
            
            # 提取元数据
            source_pdf_path = parsed_data.get("meta", {}).get("source_file", "未知来源")
            doc_title = parsed_data.get("meta", {}).get("title", os.path.basename(folder_path))

            # 准备要存入的文本块和元数据
            chunks = []
            metadatas = []
            
            # 使用sections作为文本块
            for section in parsed_data.get("sections", []):
                content = section.get("content")
                if not content:
                    continue
                
                chunks.append(content)
                
                # 为每个块准备元数据
                chunk_metadata = {
                    "source_file": source_pdf_path,
                    "document_title": doc_title,
                    "project_name": project_name,
                    "section_title": section.get("title", ""),
                    "source_page": section.get("source_page", 0)
                }
                metadatas.append(chunk_metadata)

            if not chunks:
                return "✅ 文件夹处理完成，没有找到可供embedding的文本内容。"

            # 批量添加到向量库
            total_chunks_added = 0
            doc_id_prefix = hashlib.md5(source_pdf_path.encode()).hexdigest()

            for i, chunk in enumerate(chunks):
                doc_id = f"{doc_id_prefix}_section_{i}"
                self.vector_store.collection.add(
                    ids=[doc_id],
                    documents=[chunk],
                    metadatas=[metadatas[i]]
                )
                total_chunks_added += 1

            return f"✅ 成功处理文件夹 '{folder_path}'，已添加 {total_chunks_added} 个文本块到知识库，来源: {os.path.basename(source_pdf_path)}."

        except Exception as e:
            logger.error(f"处理解析文件夹失败: {e}", exc_info=True)
            return f"❌ 处理解析文件夹失败: {str(e)}"
    
    def _extract_text_from_parsed_data(self, parsed_data: Dict[str, Any]) -> str:
        """从parsed_content.json中提取所有文本内容（保留备用）"""
        text_parts = []
        
        try:
            # 处理不同的解析数据结构
            if isinstance(parsed_data, dict):
                # 递归提取所有文本内容
                self._extract_text_recursive(parsed_data, text_parts)
            elif isinstance(parsed_data, list):
                for item in parsed_data:
                    if isinstance(item, dict):
                        self._extract_text_recursive(item, text_parts)
                    elif isinstance(item, str):
                        text_parts.append(item)
            elif isinstance(parsed_data, str):
                text_parts.append(parsed_data)
            
            # 合并所有文本
            full_text = "\n\n".join(text_parts)
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            return ""
    
    def _extract_text_recursive(self, data: Dict[str, Any], text_parts: List[str]):
        """
        递归提取字典中的文本内容
        """
        for key, value in data.items():
            if isinstance(value, str) and len(value.strip()) > 0:
                # 添加有意义的文本内容
                if len(value.strip()) > 10:  # 过滤太短的文本
                    text_parts.append(f"[{key}]: {value.strip()}")
            elif isinstance(value, dict):
                # 递归处理嵌套字典
                self._extract_text_recursive(value, text_parts)
            elif isinstance(value, list):
                # 处理列表
                for i, item in enumerate(value):
                    if isinstance(item, str) and len(item.strip()) > 10:
                        text_parts.append(f"[{key}_{i}]: {item.strip()}")
                    elif isinstance(item, dict):
                        self._extract_text_recursive(item, text_parts) 